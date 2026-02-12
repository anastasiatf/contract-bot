
import os
import asyncio
from datetime import datetime
from aiogram import Bot, Dispatcher, F
from aiogram.types import Message, ReplyKeyboardMarkup, KeyboardButton, ReplyKeyboardRemove
from aiogram.fsm.storage.memory import MemoryStorage
from aiogram.fsm.context import FSMContext
from aiogram.fsm.state import StatesGroup, State
from docx import Document

TOKEN = os.getenv("BOT_TOKEN")

bot = Bot(token=TOKEN)
dp = Dispatcher(storage=MemoryStorage())


# ================== STATES ==================

class Form(StatesGroup):
    manager = State()
    company = State()
    real_company = State()
    address = State()
    city = State()
    contract_company = State()
    reservation = State()
    stock = State()
    contract_type = State()
    delivery_time = State()
    price = State()
    extra = State()
    extra_price = State()
    payment = State()
    warranty = State()
    warranty_months = State()
    warranty_who = State()
    warranty_start = State()
    need_delivery = State()
    delivery_price = State()
    need_pnr = State()
    pnr_price = State()
    pnr_term = State()
    engineer_term = State()
    pnr_address = State()
    final_price = State()
    purchaser = State()
    brand_manager = State()
    boss = State()
    special = State()


# ================== START ==================

@dp.message(F.text == "/start")
async def start(message: Message, state: FSMContext):
    await state.clear()
    await message.answer("Введите имя и фамилию менеджера:")
    await state.set_state(Form.manager)


@dp.message(Form.manager)
async def step_manager(message: Message, state: FSMContext):
    await state.update_data(manager=message.text)
    await message.answer("Введите название компании-покупателя:")
    await state.set_state(Form.company)


@dp.message(Form.company)
async def step_company(message: Message, state: FSMContext):
    await state.update_data(company=message.text)
    await message.answer("Введите название реального покупателя (или напишите 'Совпадает'):")
    await state.set_state(Form.real_company)


@dp.message(Form.real_company)
async def step_real_company(message: Message, state: FSMContext):
    await state.update_data(real_company=message.text)
    await message.answer("Введите адрес доставки:")
    await state.set_state(Form.address)


@dp.message(Form.address)
async def step_address(message: Message, state: FSMContext):
    await state.update_data(address=message.text)
    await message.answer("Введите город поставки:")
    await state.set_state(Form.city)


@dp.message(Form.city)
async def step_city(message: Message, state: FSMContext):
    await state.update_data(city=message.text)
    keyboard = ReplyKeyboardMarkup(
        keyboard=[
            [KeyboardButton(text="Майхонг Трейдинг")],
            [KeyboardButton(text="Лимао Трейдинг")],
            [KeyboardButton(text="МирОснастки")]
        ],
        resize_keyboard=True
    )
    await message.answer("От какой компании делаем договор?", reply_markup=keyboard)
    await state.set_state(Form.contract_company)


@dp.message(Form.contract_company)
async def step_contract_company(message: Message, state: FSMContext):
    await state.update_data(contract_company=message.text)
    await message.answer("Введите номер брони в 1С (или напишите 'Нет'):",
                         reply_markup=ReplyKeyboardRemove())
    await state.set_state(Form.reservation)


@dp.message(Form.reservation)
async def step_reservation(message: Message, state: FSMContext):
    await state.update_data(reservation=message.text)
    await message.answer("Товар в наличии? (Да / В пути / Заказная позиция)")
    await state.set_state(Form.stock)


@dp.message(Form.stock)
async def step_stock(message: Message, state: FSMContext):
    await state.update_data(stock=message.text)
    await message.answer("Вид договора (Поставка / Счет-оферта / ПНР):")
    await state.set_state(Form.contract_type)


@dp.message(Form.contract_type)
async def step_contract_type(message: Message, state: FSMContext):
    await state.update_data(contract_type=message.text)
    await message.answer("Срок поставки в днях (или 'Предоплата не предусмотрена'):")
    await state.set_state(Form.delivery_time)


@dp.message(Form.delivery_time)
async def step_delivery_time(message: Message, state: FSMContext):
    await state.update_data(delivery_time=message.text)
    await message.answer("Стоимость товара по прайсу:")
    await state.set_state(Form.price)


@dp.message(Form.price)
async def step_price(message: Message, state: FSMContext):
    await state.update_data(price=message.text)
    await message.answer("Есть ли дополнительная комплектация? (Да / Нет)")
    await state.set_state(Form.extra)


@dp.message(Form.extra)
async def step_extra(message: Message, state: FSMContext):
    await state.update_data(extra=message.text)
    if message.text.lower() == "да":
        await message.answer("Введите стоимость дополнительной комплектации:")
        await state.set_state(Form.extra_price)
    else:
        await state.set_state(Form.payment)
        await message.answer("Введите условия оплаты:")


@dp.message(Form.extra_price)
async def step_extra_price(message: Message, state: FSMContext):
    await state.update_data(extra_price=message.text)
    await message.answer("Введите условия оплаты:")
    await state.set_state(Form.payment)


@dp.message(Form.payment)
async def step_payment(message: Message, state: FSMContext):
    await state.update_data(payment=message.text)
    await message.answer("Есть ли гарантия? (Да / Нет)")
    await state.set_state(Form.warranty)


@dp.message(Form.warranty)
async def step_warranty(message: Message, state: FSMContext):
    await state.update_data(warranty=message.text)
    if message.text.lower() == "да":
        await message.answer("Количество месяцев гарантии:")
        await state.set_state(Form.warranty_months)
    else:
        await state.set_state(Form.need_delivery)
        await message.answer("Требуется ли доставка? (Да / Нет)")


@dp.message(Form.warranty_months)
async def step_warranty_months(message: Message, state: FSMContext):
    await state.update_data(warranty_months=message.text)
    await message.answer("Кто несет гарантию? (Наша компания / Завод-производитель)")
    await state.set_state(Form.warranty_who)


@dp.message(Form.warranty_who)
async def step_warranty_who(message: Message, state: FSMContext):
    await state.update_data(warranty_who=message.text)
    await message.answer("С какого момента начинается гарантия?")
    await state.set_state(Form.warranty_start)


@dp.message(Form.warranty_start)
async def step_warranty_start(message: Message, state: FSMContext):
    await state.update_data(warranty_start=message.text)
    await message.answer("Требуется ли доставка? (Да / Нет)")
    await state.set_state(Form.need_delivery)


@dp.message(Form.need_delivery)
async def step_need_delivery(message: Message, state: FSMContext):
    await state.update_data(need_delivery=message.text)
    if message.text.lower() == "да":
        await message.answer("Плановая стоимость доставки:")
        await state.set_state(Form.delivery_price)
    else:
        await state.set_state(Form.need_pnr)
        await message.answer("Требуется ли ПНР? (Да / Нет)")


@dp.message(Form.delivery_price)
async def step_delivery_price(message: Message, state: FSMContext):
    await state.update_data(delivery_price=message.text)
    await message.answer("Требуется ли ПНР? (Да / Нет)")
    await state.set_state(Form.need_pnr)


@dp.message(Form.need_pnr)
async def step_need_pnr(message: Message, state: FSMContext):
    await state.update_data(need_pnr=message.text)
    if message.text.lower() == "да":
        await message.answer("Плановая стоимость ПНР:")
        await state.set_state(Form.pnr_price)
    else:
        await state.set_state(Form.final_price)
        await message.answer("Итоговая стоимость для спецификации:")


@dp.message(Form.pnr_price)
async def step_pnr_price(message: Message, state: FSMContext):
    await state.update_data(pnr_price=message.text)
    await message.answer("Срок проведения ПНР:")
    await state.set_state(Form.pnr_term)


@dp.message(Form.pnr_term)
async def step_pnr_term(message: Message, state: FSMContext):
    await state.update_data(pnr_term=message.text)
    await message.answer("Срок выезда инженера:")
    await state.set_state(Form.engineer_term)


@dp.message(Form.engineer_term)
async def step_engineer_term(message: Message, state: FSMContext):
    await state.update_data(engineer_term=message.text)
    await message.answer("Адрес проведения ПНР:")
    await state.set_state(Form.pnr_address)


@dp.message(Form.pnr_address)
async def step_pnr_address(message: Message, state: FSMContext):
    await state.update_data(pnr_address=message.text)
    await message.answer("Итоговая стоимость для спецификации:")
    await state.set_state(Form.final_price)


@dp.message(Form.final_price)
async def step_final_price(message: Message, state: FSMContext):
    await state.update_data(final_price=message.text)
    await message.answer("Есть ли закупщик? (Имя или 'Нет')")
    await state.set_state(Form.purchaser)


@dp.message(Form.purchaser)
async def step_purchaser(message: Message, state: FSMContext):
    await state.update_data(purchaser=message.text)
    await message.answer("Есть ли бренд-менеджер? (Имя или 'Нет')")
    await state.set_state(Form.brand_manager)


@dp.message(Form.brand_manager)
async def step_brand_manager(message: Message, state: FSMContext):
    await state.update_data(brand_manager=message.text)
    await message.answer("Согласующий руководитель:")
    await state.set_state(Form.boss)


@dp.message(Form.boss)
async def step_boss(message: Message, state: FSMContext):
    await state.update_data(boss=message.text)
    await message.answer("Особые условия (или 'Нет'):")
    await state.set_state(Form.special)


@dp.message(Form.special)
async def step_special(message: Message, state: FSMContext):
    await state.update_data(special=message.text)
    data = await state.get_data()

    document = Document()
    document.add_heading("ЗАЯВКА НА ДОГОВОР", level=1)
    document.add_paragraph(f"Дата: {datetime.now().strftime('%d.%m.%Y')}")

    document.add_heading("1. Общая информация", level=2)
    document.add_paragraph(f"Менеджер: {data.get('manager','')}")
    document.add_paragraph(f"Компания-покупатель: {data.get('company','')}")
    document.add_paragraph(f"Реальный покупатель: {data.get('real_company','')}")
    document.add_paragraph(f"Адрес доставки: {data.get('address','')}")
    document.add_paragraph(f"Город поставки: {data.get('city','')}")
    document.add_paragraph(f"Компания договора: {data.get('contract_company','')}")

    document.add_heading("2. Коммерческие условия", level=2)
    document.add_paragraph(f"Номер брони: {data.get('reservation','')}")
    document.add_paragraph(f"Наличие товара: {data.get('stock','')}")
    document.add_paragraph(f"Вид договора: {data.get('contract_type','')}")
    document.add_paragraph(f"Срок поставки: {data.get('delivery_time','')}")
    document.add_paragraph(f"Стоимость по прайсу: {data.get('price','')}")

    document.add_heading("3. Финансовый итог", level=2)
    document.add_paragraph(f"Итоговая стоимость: {data.get('final_price','')}")

    document.add_heading("4. Ответственные лица", level=2)
    document.add_paragraph(f"Закупщик: {data.get('purchaser','')}")
    document.add_paragraph(f"Бренд-менеджер: {data.get('brand_manager','')}")
    document.add_paragraph(f"Согласующий руководитель: {data.get('boss','')}")

    document.add_heading("5. Особые условия", level=2)
    document.add_paragraph(data.get('special',''))

    file_name = "zayavka.docx"
    document.save(file_name)

    await message.answer_document(open(file_name, "rb"))
    await state.clear()


async def main():
    await dp.start_polling(bot)


if __name__ == "__main__":
    asyncio.run(main())
